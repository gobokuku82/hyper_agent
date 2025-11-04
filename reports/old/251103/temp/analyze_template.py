"""주택임대차 표준계약서 템플릿 구조 분석 스크립트"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from pathlib import Path

# 템플릿 경로
template_path = Path(__file__).parent / "data" / "storage" / "documents" / "주택임대차 표준계약서.docx"

doc = Document(str(template_path))

print("=" * 80)
print("주택임대차 표준계약서 템플릿 구조 분석")
print("=" * 80)

print(f"\n총 단락 수: {len(doc.paragraphs)}")
print(f"총 테이블 수: {len(doc.tables)}\n")

# 각 테이블 분석
for table_idx, table in enumerate(doc.tables):
    print(f"\n{'='*80}")
    print(f"테이블 {table_idx}: {len(table.rows)}행 × {len(table.columns)}열")
    print(f"{'='*80}")

    for row_idx, row in enumerate(table.rows):
        print(f"\n--- 행 {row_idx} (셀 수: {len(row.cells)}) ---")

        for col_idx, cell in enumerate(row.cells):
            text = cell.text.strip()
            if text:  # 텍스트가 있는 셀만 출력
                # 긴 텍스트는 잘라서 표시
                display_text = text[:80] if len(text) > 80 else text
                display_text = display_text.replace('\n', '\\n')
                print(f"  [{row_idx},{col_idx}]: {display_text}")

    # 테이블 간격
    if table_idx < len(doc.tables) - 1:
        print()
