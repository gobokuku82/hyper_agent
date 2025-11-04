"""원본 템플릿에 플레이스홀더 추가 스크립트"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from pathlib import Path
from datetime import datetime

# 템플릿 경로
template_path = Path(__file__).parent / "data" / "storage" / "documents" / "주택임대차 표준계약서.docx"
output_path = template_path.parent / "주택임대차 표준계약서_with_placeholders.docx"

doc = Document(str(template_path))

print("=" * 80)
print("주택임대차 표준계약서 템플릿에 플레이스홀더 추가")
print("=" * 80)

# 테이블 0 - 주요 계약 정보
table0 = doc.tables[0]

# 행 2, 열 2-7: 소재지 (도로명주소)
for col_idx in range(2, 8):
    if col_idx < len(table0.rows[2].cells):
        table0.rows[2].cells[col_idx].text = "{{address_road}}"
print("✓ 소재지 (도로명주소) 플레이스홀더 추가")

# 행 5, 열 2-4: 임차할 부분 (상세주소)
for col_idx in range(2, 5):
    if col_idx < len(table0.rows[5].cells):
        table0.rows[5].cells[col_idx].text = "{{address_detail}}"
print("✓ 임차할 부분 (상세주소) 플레이스홀더 추가")

# 행 5, 열 5: 면적
if 5 < len(table0.rows[5].cells):
    table0.rows[5].cells[5].text = "{{rental_area}}"
print("✓ 임차 면적 플레이스홀더 추가")

# 행 11: 보증금
# 열 1-7 병합된 셀
table0.rows[11].cells[1].text = "금 {{deposit_hangeul}} 원정(₩{{deposit}})"
print("✓ 보증금 플레이스홀더 추가")

# 행 12: 계약금
table0.rows[12].cells[1].text = "금 {{contract_payment}} 원정(₩{{contract_payment}})은 계약시에 지불하고 영수함. 영수자 (         인)"
print("✓ 계약금 플레이스홀더 추가")

# 행 15: 차임(월세)
table0.rows[15].cells[1].text = "금 {{monthly_rent}} 원정은  매월 {{monthly_rent_day}}일에 지불한다(입금계좌:                                  )"
print("✓ 월세 플레이스홀더 추가")

# 행 16: 관리비
table0.rows[16].cells[1].text = "(정액인 경우) 금 {{management_fee}} 원정(₩{{management_fee}})"
print("✓ 관리비 플레이스홀더 추가")

# 테이블 4 - 당사자 정보 (11행 × 8열이지만 병합으로 실제 셀 수는 15개)
table4 = doc.tables[4]

# 임대인 정보
# 행 0: 주소
if 2 < len(table4.rows[0].cells):
    # 주소 셀 찾기 (병합된 셀이므로 큰 셀에 입력)
    for cell_idx, cell in enumerate(table4.rows[0].cells):
        if "주        소" in cell.text or cell_idx == 2:
            # 주소 입력 셀 (보통 큰 병합 셀)
            table4.rows[0].cells[2].text = "{{lessor_address}}"
            break
print("✓ 임대인 주소 플레이스홀더 추가")

# 행 1: 성명
if 9 < len(table4.rows[1].cells):
    table4.rows[1].cells[9].text = "{{lessor_name}}"
print("✓ 임대인 성명 플레이스홀더 추가")

# 행 1: 전화
if 6 < len(table4.rows[1].cells):
    # 전화번호 입력 셀 찾기
    for cell_idx in range(6, min(9, len(table4.rows[1].cells))):
        if "전        화" not in table4.rows[1].cells[cell_idx].text:
            table4.rows[1].cells[cell_idx].text = "{{lessor_phone}}"
            break
print("✓ 임대인 전화 플레이스홀더 추가")

# 임차인 정보
# 행 3: 주소
if 2 < len(table4.rows[3].cells):
    table4.rows[3].cells[2].text = "{{lessee_address}}"
print("✓ 임차인 주소 플레이스홀더 추가")

# 행 4: 성명
if 9 < len(table4.rows[4].cells):
    table4.rows[4].cells[9].text = "{{lessee_name}}"
print("✓ 임차인 성명 플레이스홀더 추가")

# 행 4: 전화
if 6 < len(table4.rows[4].cells):
    for cell_idx in range(6, min(9, len(table4.rows[4].cells))):
        if "전        화" not in table4.rows[4].cells[cell_idx].text:
            table4.rows[4].cells[cell_idx].text = "{{lessee_phone}}"
            break
print("✓ 임차인 전화 플레이스홀더 추가")

# 저장
doc.save(str(output_path))

print(f"\n{'='*80}")
print(f"✅ 플레이스홀더가 추가된 템플릿 저장 완료")
print(f"📁 저장 경로: {output_path}")
print(f"{'='*80}")
